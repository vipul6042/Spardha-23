from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from io import BytesIO

def create_form(data):

    # Create a new Word document
    doc = Document()

    # Main Heading (SPARDHA'23)
    main_heading = doc.add_heading("SPARDHA'23", level=0)
    main_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    main_heading.runs[0].font.size = Pt(36)

    # Sub Heading (DETAILED ENTRY FORM)
    sub_heading = doc.add_heading('DETAILED ENTRY FORM', level=1)
    sub_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_heading.runs[0].font.size = Pt(15)

    # College Name
    college_name = doc.add_paragraph(f"\t\tName of college: {data['institution_name']}")
    college_name.runs[0].font.size = Pt(18)

    initial_headings = (
        ('Total No. of Participants', f'Female: {data["num_of_girls"]}', f'Male: {data["num_of_girls"]}'),
        ('No. of officials', f'Faculty Members: {data["num_of_faculty_members"]}', f'Coaches & PTI: {data["num_of_coaches_pti"]}', f'Supporting Staff: {data["num_of_faculty_members"]}'),
        ('Arrival of the Team:','Date','Train/ Flight No.:','PNR nos. (If via Train):'),
        ('Departure of the Team:','Train/ Flight No.:','PNR nos. (If via Train):'),
    )

    i = 1
    for tupl in initial_headings:
        heading = doc.add_heading(f'{i}. {tupl[0]}\n', level=2)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading.runs[0].bold = True    
        for text in tupl[1:]:
            para = doc.add_paragraph(f"{text}")
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        i += 1

    contact_info = ("Sports Officer", "Sports Secretary")

    heading = doc.add_heading(f'{i}. Contact Information of Contingent Leaders:', level=2)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    heading.runs[0].bold = True

    j = 1
    for heading in contact_info:
        para = doc.add_heading(f"\t{j}) {heading}:\n", level=3)
        doc.add_paragraph("\tName:")
        doc.add_paragraph("\tPh No.:")
        j += 1

    return_info = doc.add_paragraph("\nKindly return the Performa, duly completed in all respect, and send it by email at:")
    email_address = doc.add_paragraph("Email: convener.spardha@iitbhu.ac.in")
    email_address.runs[0].bold = True
    instructions = doc.add_paragraph("Also, you have to bring 2 hardcopies of this form which should be completely filled and duly signed by dean/director of your college at the time of registration.")

    signature_heading = doc.add_paragraph('\n\n\n\nSignature with Seal:')
    signature_heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    note = doc.add_paragraph("\n\nNote: Name of the Participants to be given in full and in block letters.")
    note.runs[0].bold = True

    def make_row_bold(row):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    for game in data["games"]:
        game_type = "Women" if game["game_type"] == "G" else ("Men" if game["game_type"] == "B" else "Mixed")
        heading = doc.add_heading(f'{game["game_name"].upper()} ({game_type})\n', level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading.runs[0].bold = True

        table = doc.add_table(rows=game["max_players"]+1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        for row in table.rows:
            row.height = Inches(0.4)

        for cell in table.columns[0].cells:
            cell.width = Inches(1)
        for cell in table.columns[1].cells:
            cell.width = Inches(5)
        p = table.rows[0].cells[0].add_paragraph("S.No.")
        p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
        p = table.rows[0].cells[1].add_paragraph("Name Of the Player")
        p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
        make_row_bold(table.rows[0])
        i = 1
        for player in game["players"]:
            cell_row = table.rows[i].cells
            p = cell_row[0].add_paragraph(str(i))
            p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell_row[1].add_paragraph(str(player))
            p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            i += 1

        
        captain_info = doc.add_paragraph(f"\n\nCaptain Name: {game['captain_name']}     Email ID: ________________________________________\n\nPh.no. : {game['captain_phone']}")
        captain_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        vice_captain_info = doc.add_paragraph("Vice-Captain Name: _________________________ Email ID: ________________________________________\n\nPh.no.:  ______________________")
        vice_captain_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        signature_heading = doc.add_paragraph('\n\n\n\nSignature with Seal')
        signature_heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
